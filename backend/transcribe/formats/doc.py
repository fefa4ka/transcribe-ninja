#!/usr/bin/python
# -*- coding: utf-8 -*-
from django.conf import settings

import os

import time

import codecs

from hashlib import md5

from docx import Document
from docx.shared import Inches


class ExportDOC:
    def _to_doc(self, transcriptions, file_name):
        # file_name = md5(
        #     str(time.time())
        # ).hexdigest()
        doc_file_path = os.path.join(settings.TEMP_DIR, "%s.doc" % file_name)

        document = Document()

        document.add_heading(u'Стенограмма записи «%s»' % self.record.title, 0)

        document.add_paragraph(u'Расшифровал Стеня Графов\nСтенограф.ус\nhttp://stenograph.us')

        for index, transcription in enumerate(transcriptions):
            start_at = time.strftime(
                u"%H:%M:%S",
                time.gmtime(transcription.start_at)
            )

            text = document.add_paragraph('%s \n' % start_at)
            text.add_run('%s: ' % transcription.export_name).bold = True
            text.add_run(transcription.text_pretty)

        document.add_page_break()
        document.save(doc_file_path)

        return doc_file_path
